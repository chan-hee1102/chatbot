"""Convert docs/*.md to polished docs/*.docx via python-docx.

Usage (from project root):
    python scripts/build_docs.py

Designed for the two top-level documents in this repo:
  - docs/01_서비스기획안.md
  - docs/02_요구사항명세서.md

Handles: headings (H1-H4), paragraphs, bullet/numbered lists,
GFM-style pipe tables, fenced code blocks, inline `code` and **bold**.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = PROJECT_ROOT / "docs"
SOURCES = ["01_서비스기획안.md", "02_요구사항명세서.md"]

BODY_FONT = "맑은 고딕"
CODE_FONT = "Consolas"
HEADING_COLOR = RGBColor(0x11, 0x18, 0x27)  # gray-900
META_COLOR = RGBColor(0x6B, 0x72, 0x80)     # gray-500
RULE_COLOR = "D1D5DB"                       # gray-300

# --- Low-level styling helpers --------------------------------------------------


def _set_eastasia_font(run, name: str) -> None:
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def _style_run(run, *, font=BODY_FONT, size=10.5, bold=False,
               italic=False, color=None) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    _set_eastasia_font(run, font)


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_bottom_rule(paragraph) -> None:
    """Add a thin bottom border to act as a section divider."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RULE_COLOR)
    pBdr.append(bottom)
    pPr.append(pBdr)


# --- Inline markdown rendering --------------------------------------------------

INLINE_PATTERN = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")


def _render_inline(paragraph, text: str, *, base_size: float = 10.5,
                   base_color=None) -> None:
    """Render a single line with inline `code` and **bold** support."""
    pos = 0
    for match in INLINE_PATTERN.finditer(text):
        start, end = match.span()
        if start > pos:
            run = paragraph.add_run(text[pos:start])
            _style_run(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _style_run(run, font=CODE_FONT, size=base_size - 0.5,
                       color=RGBColor(0x11, 0x18, 0x27))
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F3F4F6")
            run._element.get_or_add_rPr().append(shd)
        else:  # **bold**
            run = paragraph.add_run(token[2:-2])
            _style_run(run, size=base_size, bold=True, color=base_color)
        pos = end
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _style_run(run, size=base_size, color=base_color)


# --- Block-level builders ------------------------------------------------------


def _add_heading(doc: Document, level: int, text: str) -> None:
    sizes = {1: 22, 2: 16, 3: 13, 4: 11.5}
    size = sizes.get(level, 11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _style_run(run, size=size, bold=True, color=HEADING_COLOR)
    if level == 1:
        _add_bottom_rule(p)


def _add_paragraph(doc: Document, text: str, *, color=None,
                   size: float = 10.5, align=None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if align is not None:
        p.alignment = align
    _render_inline(p, text, base_size=size, base_color=color)


def _add_list_item(doc: Document, text: str, *, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    _render_inline(p, text)
    for run in p.runs:
        _style_run(run, size=10.5)


def _add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run("\n".join(lines))
    _style_run(run, font=CODE_FONT, size=9.5,
               color=RGBColor(0x11, 0x18, 0x27))
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F9FAFB")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "E5E7EB")
        pBdr.append(el)
    pPr.append(pBdr)


def _parse_table_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _render_inline(p, cell_text, base_size=9.5)
            if r_idx == 0:
                _shade_cell(cell, "F3F4F6")
                for run in p.runs:
                    run.font.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


# --- Markdown stream parser ---------------------------------------------------


def _is_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\|?[\s:|-]+\|?", line.strip())) and "-" in line


def _build_docx(md_path: Path, out_path: Path) -> None:
    doc = Document()

    # Base style tweaks
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    _set_eastasia_font(style.element.rPr.rFonts if False else style.font, BODY_FONT)  # no-op safeguard
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), BODY_FONT)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped in {"---", "***", "___"}:
            p = doc.add_paragraph()
            _add_bottom_rule(p)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            _add_heading(doc, level, m.group(2).strip())
            i += 1
            continue

        # Fenced code blocks
        if stripped.startswith("```"):
            i += 1
            block: list[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # consume closing fence
            _add_code_block(doc, block)
            continue

        # Tables: header line then separator line then body
        if stripped.startswith("|") and i + 1 < n and _is_separator(lines[i + 1]):
            rows: list[list[str]] = [_parse_table_row(stripped)]
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_parse_table_row(lines[i].strip()))
                i += 1
            _add_table(doc, rows)
            continue

        # Lists
        if re.match(r"^\s*[-*]\s+", line):
            _add_list_item(doc, re.sub(r"^\s*[-*]\s+", "", line), numbered=False)
            i += 1
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            _add_list_item(doc, re.sub(r"^\s*\d+\.\s+", "", line), numbered=True)
            i += 1
            continue

        # Front matter style metadata (lines starting with **)
        if stripped.startswith("**") and "**" in stripped[2:]:
            _add_paragraph(doc, stripped, size=10.5)
            i += 1
            continue

        # Default: paragraph
        _add_paragraph(doc, stripped)
        i += 1

    doc.save(out_path)
    print(f"  → wrote {out_path.relative_to(PROJECT_ROOT)}")


def main() -> None:
    for name in SOURCES:
        md = DOCS_DIR / name
        out = DOCS_DIR / (Path(name).stem + ".docx")
        print(f"Building {name} …")
        _build_docx(md, out)


if __name__ == "__main__":
    main()
